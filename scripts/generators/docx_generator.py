"""Generate synthetic Word (.docx) test files using python-docx."""

from pathlib import Path

from generators.base import BaseGenerator

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def _require_docx() -> None:
    if not _DOCX_OK:
        raise ImportError("python-docx is required: uv add python-docx")


class DocxGenerator(BaseGenerator):
    category = "docx"

    def generate(self, dry_run: bool = False) -> list[dict]:
        _require_docx()
        out = self.output_dir()
        if not dry_run:
            out.mkdir(parents=True, exist_ok=True)

        entries = []
        entries.append(self._simple_doc(out, dry_run))
        entries.append(self._tables_merged(out, dry_run))
        entries.append(self._nested_tables(out, dry_run))
        entries.append(self._inline_images(out, dry_run))
        entries.append(self._headers_footers(out, dry_run))
        return entries

    # ------------------------------------------------------------------
    def _simple_doc(self, out: Path, dry_run: bool) -> dict:
        filename = "syn_simple_doc.docx"
        if dry_run:
            print(f"  [dry-run] would create {self.category}/{filename}")
        else:
            doc = Document()
            doc.add_heading("Project Proposal: Data Analytics Platform", level=1)
            doc.add_paragraph(
                "Executive Summary\n"
                "This proposal outlines the development of a centralized data analytics "
                "platform to consolidate reporting across all business units. The platform "
                "will provide real-time dashboards, self-service analytics, and automated "
                "reporting capabilities."
            )
            doc.add_heading("Objectives", level=2)
            for obj in [
                "Unify data from 12 disparate systems into a single source of truth",
                "Reduce report generation time from 4 hours to under 5 minutes",
                "Enable business users to build dashboards without IT involvement",
                "Provide audit trails and data lineage for compliance purposes",
            ]:
                doc.add_paragraph(obj, style="List Bullet")

            doc.add_heading("Timeline", level=2)
            doc.add_paragraph(
                "The project is planned over three phases spanning 18 months. Phase 1 "
                "(months 1-6) covers data infrastructure setup. Phase 2 (months 7-12) "
                "delivers core analytics features. Phase 3 (months 13-18) includes "
                "advanced ML-powered insights."
            )
            doc.add_heading("Budget", level=2)
            doc.add_paragraph(
                "Total estimated investment: $1.2M. This includes infrastructure costs "
                "($400K), software licensing ($250K), professional services ($350K), and "
                "internal staffing allocation ($200K)."
            )
            doc.save(out / filename)
        return self._entry(filename, "easy",
                           ["docx", "simple", "headings", "paragraphs", "bullet-list"],
                           False, "Simple Word document with headings, paragraphs, and a bullet list.")

    def _tables_merged(self, out: Path, dry_run: bool) -> dict:
        filename = "syn_tables_merged.docx"
        if dry_run:
            print(f"  [dry-run] would create {self.category}/{filename}")
        else:
            doc = Document()
            doc.add_heading("Quarterly Expense Report — Q3 2024", level=1)
            doc.add_paragraph("Summary of departmental expenses for Q3 2024.")

            # Table with merged header cells
            table = doc.add_table(rows=8, cols=5)
            table.style = "Table Grid"

            # Top header row: span "Quarterly Amounts" across cols 2-5
            hdr = table.rows[0].cells
            hdr[0].text = "Department"
            hdr[1].text = "July"
            hdr[2].text = "August"
            hdr[3].text = "September"
            hdr[4].text = "Q3 Total"
            for cell in hdr:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Merge the quarter-label row
            row1 = table.add_row().cells
            row1[0].text = ""
            merged = row1[1].merge(row1[4])
            merged.text = "← Quarterly Amounts →"

            data = [
                ("Engineering", 185000, 192000, 178000),
                ("Marketing", 72000, 88000, 95000),
                ("Sales", 141000, 155000, 168000),
                ("Operations", 98000, 101000, 103000),
                ("HR & Admin", 45000, 46000, 47000),
                ("Finance", 38000, 39000, 41000),
            ]
            for dept, jul, aug, sep in data:
                row = table.add_row().cells
                row[0].text = dept
                row[1].text = f"${jul:,}"
                row[2].text = f"${aug:,}"
                row[3].text = f"${sep:,}"
                row[4].text = f"${jul + aug + sep:,}"

            doc.add_paragraph("")
            doc.add_paragraph("Note: All figures in USD. Excludes capital expenditures.")
            doc.save(out / filename)
        return self._entry(filename, "medium",
                           ["docx", "table", "merged-cells", "expense-report"],
                           False, "Word document with a table containing merged header cells.")

    def _nested_tables(self, out: Path, dry_run: bool) -> dict:
        filename = "syn_nested_tables.docx"
        if dry_run:
            print(f"  [dry-run] would create {self.category}/{filename}")
        else:
            doc = Document()
            doc.add_heading("Team Directory", level=1)
            doc.add_paragraph("The following table lists each team with their members and roles.")

            outer = doc.add_table(rows=1, cols=3)
            outer.style = "Table Grid"
            hdr = outer.rows[0].cells
            hdr[0].text = "Team"
            hdr[1].text = "Manager"
            hdr[2].text = "Members"
            for cell in hdr:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            teams = [
                ("Backend", "Alice Chen", [("Bob Kim", "Sr. Engineer"), ("Carol Patel", "Engineer"), ("Dave Osei", "Jr. Engineer")]),
                ("Frontend", "Eve Torres", [("Frank Liu", "Sr. Engineer"), ("Grace Park", "Designer")]),
                ("DevOps", "Henry Bravo", [("Iris Nguyen", "Platform Engineer"), ("Jack Müller", "SRE")]),
            ]

            for team_name, manager, members in teams:
                row = outer.add_row().cells
                row[0].text = team_name
                row[1].text = manager

                # Nested table inside the Members cell
                inner = row[2].add_table(rows=1, cols=2)
                inner.style = "Table Grid"
                inner.rows[0].cells[0].text = "Name"
                inner.rows[0].cells[1].text = "Role"
                for name, role in members:
                    inner_row = inner.add_row().cells
                    inner_row[0].text = name
                    inner_row[1].text = role

            doc.save(out / filename)
        return self._entry(filename, "hard",
                           ["docx", "nested-tables", "table-in-cell", "team-directory"],
                           False, "Word document with a nested table (table embedded inside a table cell).")

    def _inline_images(self, out: Path, dry_run: bool) -> dict:
        filename = "syn_inline_images.docx"
        if dry_run:
            print(f"  [dry-run] would create {self.category}/{filename}")
        else:
            # We'll generate small placeholder images with Pillow inline
            try:
                from PIL import Image as PILImage, ImageDraw
                import io

                def make_placeholder(label: str, color: tuple, w: int = 400, h: int = 200) -> io.BytesIO:
                    img = PILImage.new("RGB", (w, h), color=color)
                    draw = ImageDraw.Draw(img)
                    draw.text((20, h // 2 - 15), label, fill=(255, 255, 255))
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    buf.seek(0)
                    return buf

                doc = Document()
                doc.add_heading("Annual Wildlife Conservation Report 2024", level=1)
                doc.add_paragraph(
                    "This report summarizes key findings from our field research conducted "
                    "across three ecosystems: temperate forest, coastal wetlands, and alpine meadow."
                )
                doc.add_heading("Temperate Forest Ecosystem", level=2)
                doc.add_paragraph(
                    "Population surveys indicate a 12% increase in bird diversity compared to 2023. "
                    "The reintroduction program for red-tailed hawks showed positive early results."
                )
                buf1 = make_placeholder("[Figure 1: Temperate Forest Survey Map]", (60, 100, 60))
                doc.add_picture(buf1, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("Figure 1: Survey grid used for bird population counts (May–August 2024).")

                doc.add_heading("Coastal Wetlands Ecosystem", level=2)
                doc.add_paragraph(
                    "Water quality measurements improved significantly following the 2023 "
                    "restoration project. Migratory waterfowl counts increased by 23%."
                )
                buf2 = make_placeholder("[Figure 2: Wetlands Water Quality Chart]", (30, 80, 130))
                doc.add_picture(buf2, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("Figure 2: Monthly water quality index (Jan–Dec 2024).")

                doc.add_heading("Alpine Meadow Ecosystem", level=2)
                doc.add_paragraph(
                    "Pollinator surveys documented 34 bee species, including 3 previously "
                    "undocumented in this region. Wildflower diversity remains stable."
                )
                buf3 = make_placeholder("[Figure 3: Alpine Species Distribution]", (150, 80, 30))
                doc.add_picture(buf3, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph("Figure 3: Pollinator species distribution across elevation zones.")

            except ImportError:
                # Fallback: create doc without images if Pillow unavailable
                doc = Document()
                doc.add_heading("Annual Wildlife Conservation Report 2024", level=1)
                doc.add_paragraph("[Pillow not available — image placeholders omitted]")

            doc.save(out / filename)
        return self._entry(filename, "medium",
                           ["docx", "inline-images", "figures", "mixed-content"],
                           False, "Word document with inline images interspersed between text sections.")

    def _headers_footers(self, out: Path, dry_run: bool) -> dict:
        filename = "syn_headers_footers.docx"
        if dry_run:
            print(f"  [dry-run] would create {self.category}/{filename}")
        else:
            doc = Document()

            # Set up header
            section = doc.sections[0]
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = "CONFIDENTIAL — Internal Use Only"
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(180, 0, 0)

            # Set up footer
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "Acme Corp | Strategy Division | Page "
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            doc.add_heading("Strategic Plan 2025–2027", level=1)
            doc.add_paragraph(
                "This document outlines the three-year strategic direction for Acme Corporation. "
                "It has been prepared by the Strategy Division and is intended for internal use only."
            )

            sections_content = [
                ("Vision & Mission", (
                    "Our vision is to be the leading provider of sustainable industrial solutions. "
                    "Our mission is to deliver innovative, reliable, and eco-friendly products "
                    "that create lasting value for our customers and communities."
                )),
                ("Strategic Priorities", (
                    "For 2025–2027, we have identified four strategic priorities:\n"
                    "1. Market Expansion — Enter three new international markets\n"
                    "2. Product Innovation — Launch next-generation product line\n"
                    "3. Operational Excellence — Achieve 20% cost reduction\n"
                    "4. Digital Transformation — Complete ERP migration and AI adoption"
                )),
                ("Financial Targets", (
                    "Revenue target: $2.4B by end of 2027 (35% growth from 2024 baseline). "
                    "EBITDA margin target: 22% (up from current 17%). "
                    "R&D investment: 8% of annual revenue."
                )),
                ("Risk Factors", (
                    "Key risks include geopolitical uncertainty, supply chain disruptions, "
                    "regulatory changes in target markets, and talent acquisition challenges. "
                    "Mitigation plans are detailed in Appendix B."
                )),
            ]

            for heading, body in sections_content:
                doc.add_heading(heading, level=2)
                doc.add_paragraph(body)
                doc.add_page_break()

            doc.save(out / filename)
        return self._entry(filename, "medium",
                           ["docx", "headers", "footers", "page-breaks", "strategic-plan"],
                           False, "Multi-page Word document with header/footer branding and page breaks.")
